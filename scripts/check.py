# -*- coding: utf-8 -*-
"""Scan a generated .docx and write a report: heading outline, Yijing-symbol counts,
and high-frequency OCR error markers. Run after every (re)generation; the goal is to
drive every error-marker count to zero (residual hits should each be a legit word).

    python check.py <docx> [--out _chk_out.txt]

The error-marker patterns below are tuned for classical-Chinese OCR; edit them for
your own text. Output goes to a UTF-8 file (console CJK mangles).
"""
import argparse, io, re
from docx import Document

MARKERS = [
    ('empty quotes 空引号',  '“”'),
    ('empty book title 空书名', '《》'),
    ('dash-in-quotes 破折号代字', '“[—–―_＿]”'),
    ('yao/jiao 交->爻',     '(交辞|交象|交位|卦交|两交|六交|初.交|上.交)'),
    ('shi 著->蓍',          '(生著|根著|著草|著圆|探著|揲著|誓圆|此著|楪蓍)'),
    ('shi 签->筮',          '(签法|占签|卜签|签辞)'),
    ('hexagram-as-重',      '[乾坤复剥临泰否观遁姤夬颐蛊贲屯蒙睽豫师比]重[卦，。；]'),
]

if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('docx'); ap.add_argument('--out', default='_chk_out.txt')
    a = ap.parse_args()
    d = Document(a.docx); ps = [p.text for p in d.paragraphs]
    o = io.open(a.out, 'w', encoding='utf-8')

    o.write('=== headings ===\n'); nh = {}
    for p in d.paragraphs:
        s = p.style.name if p.style else ''
        if s.startswith('Heading'):
            nh[s[-1]] = nh.get(s[-1], 0) + 1
            o.write(f'{s[-1]}  {p.text}\n')
    o.write(f'counts {nh}\n\n')

    txt = '\n'.join(ps)
    hexg = sum(1 for c in txt if 0x4DC0 <= ord(c) <= 0x4DFF)
    trig = sum(1 for c in txt if 0x2630 <= ord(c) <= 0x2637)
    o.write(f'symbols: hexagram(U+4DC0-4DFF)={hexg}  trigram(U+2630-2637)={trig}\n\n')

    o.write('=== error markers (aim: every count 0, residuals must be legit) ===\n')
    for label, pat in MARKERS:
        rx = re.compile(pat); lines = []
        for i, t in enumerate(ps):
            for m in rx.finditer(t):
                a0 = max(0, m.start() - 12); b0 = min(len(t), m.end() + 12)
                lines.append(f'  [{i}] ...{t[a0:b0]}...')
        o.write(f'-- {label}: {len(lines)} --\n' + '\n'.join(lines[:20]) + ('\n' if lines else ''))
    o.close(); print('wrote', a.out)
