# -*- coding: utf-8 -*-
"""
One-shot: MinerU OCR markdown -> corrected, structured .docx

Usage:
  # use existing MinerU markdown (recommended: lets you re-run in seconds)
  python convert.py <pdf> -o <out.docx> --md <mineru_output.md> [--toc toc_xx.py]
  # or run OCR inline (needs MinerU installed at C:\\mineru-env)
  python convert.py <pdf> -o <out.docx> --lang ch_server [--toc toc_xx.py]

Pipeline inside build_docx():
  apply corrections.py  ->  restore hexagram/trigram symbols  ->
  map titles via the toc module  ->  write .docx
    (CJK body = SongTi 10.5pt with first-line indent;
     Yijing symbols get Segoe UI Symbol, otherwise SongTi renders them blank)

NOTE ON DOMAIN SPECIFICS: the hexagram/trigram logic below is for the *Yijing*
(易经) book this toolkit was built for. For a different book, delete the
boshu()/inline_sym()/is_tbl()/is_seq64() helpers and their call sites — the rest
(corrections + toc title mapping + docx writing) is general.
"""
import os, re, glob, argparse, subprocess, importlib.util
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
def _load(name, path):
    s = importlib.util.spec_from_file_location(name, path)
    m = importlib.util.module_from_spec(s); s.loader.exec_module(m); return m
corr = _load('corr', os.path.join(HERE, 'corrections.py'))

# ---- Yijing symbol tables (domain-specific) ----
NAMES = ['乾','坤','屯','蒙','需','讼','师','比','小畜','履','泰','否','同人','大有','谦','豫',
         '随','蛊','临','观','噬嗑','贲','剥','复','无妄','大畜','颐','大过','坎','离','咸','恒',
         '遁','大壮','晋','明夷','家人','睽','蹇','解','损','益','夬','姤','萃','升','困','井',
         '革','鼎','震','艮','渐','归妹','丰','旅','巽','兑','涣','节','中孚','小过','既济','未济']
H = {n: chr(0x4DC0 + i) for i, n in enumerate(NAMES)}   # 64 hexagram chars U+4DC0..
SEQ64 = '，'.join(H[n] + n for n in NAMES) + '。'
TBL = '☰乾　☷坤　☳震　☴巽　☵坎　☲离　☶艮　☱兑'
TRIG = '☰☱☲☳☴☵☶☷'
NAMEPAT = '|'.join(re.escape(n) for n in sorted(NAMES, key=len, reverse=True))
INLINE = re.compile('(' + NAMEPAT + ')(卦?)([' + TRIG + '])')

def boshu(t):   # Mawangdui silk-text passage: tolerant fixups for its mixed symbols
    t = re.sub(r'(上卦皆乾)[三王☰]?', r'\1☰', t)
    t = re.sub(r'(上卦皆艮)[三王☶]?', r'\1☶', t)
    t = re.sub(r'(上卦皆坎)[三王☵]?', r'\1☵', t)
    t = re.sub(r'(首卦为)[' + TRIG + ']?乾', r'\1' + H['乾'] + '乾', t)
    t = re.sub(r'(次卦为)[' + TRIG + ']?否', r'\1' + H['否'] + '否', t)
    t = re.sub(r'(最后一卦为)[' + TRIG + ']?益', r'\1' + H['益'] + '益', t)
    return t
def inline_sym(t):   # body text: "<hexagram-name>(卦)?☰" -> correct hexagram char
    return INLINE.sub(lambda m: m.group(1) + m.group(2) + H[m.group(1)], t)
def is_tbl(t):       # the eight-trigram table: CJK content is exactly the 8 trigram names
    return re.sub(r'[^乾坤震巽坎离艮兑]', '', t) == '乾坤震巽坎离艮兑' and len(t) < 20
def is_seq64(t):     # the King-Wen 64-hexagram sequence line (signature, char-independent)
    return ('未济。' in t and '既济' in t and '噬嗑' in t and '归妹' in t and t.count('，') > 30)

def run_mineru(pdf, lang):
    out = os.path.join(HERE, '_mineru_out')
    env = dict(os.environ, MINERU_MODEL_SOURCE='huggingface')
    exe = r'C:\mineru-env\Scripts\mineru.exe'
    subprocess.run([exe, '-p', pdf, '-o', out, '-b', 'pipeline', '-m', 'ocr',
                    '-l', lang, '-f', 'false', '-t', 'false'], env=env, check=True)
    return glob.glob(os.path.join(out, '**', '*.md'), recursive=True)[0]

def build_docx(md, out_docx, toc=None):
    raw = open(md, encoding='utf-8').read().split('\n')
    # normalise a title to a match key: drop punctuation / digits / ordinal numerals
    def sk(s): return re.sub(r'[\s　，。、；：《》「」“”·\(\)（）①-⑩]|[0-9一二三四五六七八九十]', '', s)
    SEC = (toc.SEC if toc else {}); SUB = (toc.SUB if toc else {})
    SEC = {sk(k): v for k, v in SEC.items()}; SUB = {sk(k): v for k, v in SUB.items()}
    CHAPTERS = list(getattr(toc, 'CHAPTERS', [])) if toc else []
    ins = [False] * len(CHAPTERS)
    def maybe_chapter(disp):   # insert the chapter title before its first section (OCR often drops it)
        for k, (fs, ch) in enumerate(CHAPTERS):
            if not ins[k] and disp == fs:
                blocks.append(('h1', ch)); ins[k] = True; return
    blocks = []; i = 0
    while i < len(raw):
        ln = raw[i].rstrip()
        if not ln.strip() or ln.strip().startswith('!['):
            i += 1; continue
        if ln.startswith('##'):
            txt = ln.lstrip('#').strip()
            if not txt: i += 1; continue
            key = sk(txt)
            # "第N节" alone on its line -> combine with the following ## line as the section title
            if re.match(r'^第[一二三四五六七八九十]+节$', txt):
                nxt = raw[i + 2].lstrip('#').strip() if i + 2 < len(raw) and raw[i + 2].startswith('##') else ''
                full = sk(txt + nxt)
                if full in SEC: lvl, disp = SEC[full]
                else: lvl, disp = 'h2', txt + '　' + corr.fix(nxt)
                maybe_chapter(disp)
                blocks.append((lvl, disp)); i += (4 if nxt else 2); continue
            if txt in ('人名索引', '书名索引', '事项索引'):           # index documents
                blocks.append(('h2', txt)); i += 2; continue
            if re.fullmatch(r'[一二三四五六七八九十]{1,3}画', txt):    # index radical-stroke groups
                blocks.append(('h3', txt)); i += 2; continue
            if key in SEC:
                lvl, disp = SEC[key]; maybe_chapter(disp)
                blocks.append((lvl, disp)); i += 2; continue
            if key in SUB:
                blocks.append(('h3', SUB[key])); i += 2; continue
            if re.match(r'^[\(（]\d', txt):                          # "(1) ..." sub-points
                blocks.append(('h4', corr.fix(txt))); i += 2; continue
            if re.match(r'^第[一二三四五六七八九十]+章', txt) and len(txt) <= 16:
                blocks.append(('h1', corr.fix(txt))); i += 2; continue
            blocks.append(('p', corr.fix(txt))); i += 2; continue
        t = corr.fix(ln.strip())
        if is_tbl(t): blocks.append(('tbl', TBL)); i += 1; continue
        if is_seq64(t): blocks.append(('seq', SEQ64)); i += 1; continue
        if t.startswith('马王堆汉墓出土的帛书本') or '上卦皆乾' in t:
            blocks.append(('p', boshu(t))); i += 1; continue
        blocks.append(('p', inline_sym(t))); i += 1
    # ---- write docx ----
    doc = Document()
    st = doc.styles['Normal']; st.font.name = '宋体'; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # symbols SongTi can't render: hexagrams U+4DC0-4DFF, trigrams U+2630-2637, U+2059 (⁙)
    isx = lambda c: 0x4DC0 <= ord(c) <= 0x4DFF or 0x2630 <= ord(c) <= 0x2637 or ord(c) == 0x2059
    def add(par, text):
        buf = ''; hm = None
        def fl(s, h):
            if not s: return
            r = par.add_run(s)
            if h:
                r.font.name = 'Segoe UI Symbol'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Segoe UI Symbol')
        for ch in text:
            h = isx(ch)
            if hm is None: hm = h; buf = ch
            elif h == hm: buf += ch
            else: fl(buf, hm); buf = ch; hm = h
        fl(buf, hm)
    lm = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4}
    for lvl, x in blocks:
        if lvl in lm: add(doc.add_heading(level=lm[lvl]), x)
        elif lvl == 'tbl':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add(p, x)
        elif lvl == 'seq': add(doc.add_paragraph(), x)
        else:
            p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(21); add(p, x)
    doc.save(out_docx)
    nh = sum(1 for l, _ in blocks if l in lm)
    return len(blocks), nh

if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('pdf'); ap.add_argument('-o', '--out', required=True)
    ap.add_argument('--lang', default='ch_server')
    ap.add_argument('--toc', default=None, help='title-mapping module, e.g. toc_05.py')
    ap.add_argument('--md', default=None, help='existing MinerU markdown (skips OCR)')
    a = ap.parse_args()
    toc = _load('toc', a.toc) if a.toc else None
    md = a.md or run_mineru(a.pdf, a.lang)
    nb, nh = build_docx(md, a.out, toc)
    print('OK ->', a.out, '| blocks', nb, 'headings', nh)
